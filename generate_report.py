"""生成项目完成报告"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# 创建文档
doc = Document()


# 设置中文字体
def set_chinese_font(run, font_name="宋体", font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)


# 标题
title = doc.add_heading("", level=0)
run = title.add_run("明禾陪伴（Minghe Companion）")
run.font.name = "黑体"
run.font.size = Pt(22)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
run = subtitle.add_run("心理资讯大师 AI Agent 项目完成报告")
run.font.name = "宋体"
run.font.size = Pt(14)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 项目概述
h1 = doc.add_heading("一、项目概述", level=1)
p = doc.add_paragraph()
run = p.add_run(
    "明禾陪伴是一款基于人工智能的心理健康服务 Agent，旨在为用户提供全生命周期的心理陪伴与咨询服务。"
)
run.font.size = Pt(12)
p = doc.add_paragraph()
run = p.add_run(
    "项目采用 Python + LangChain/LangGraph 技术栈，集成 DeepSeek 大语言模型，为用户带来智能、温暖的心理健康服务体验。"
)
run.font.size = Pt(12)

doc.add_paragraph()

# 技术架构
h1 = doc.add_heading("二、技术架构", level=1)

p = doc.add_paragraph()
run = p.add_run("核心技术栈：")
run.bold = True
run.font.size = Pt(12)
run = p.add_run(" Python 3.10+ | LangChain | LangGraph | DeepSeek Chat")

p = doc.add_paragraph()
run = p.add_run("大语言模型：")
run.bold = True
run.font.size = Pt(12)
run = p.add_run(" DeepSeek-chat (deepseek-chat)")

p = doc.add_paragraph()
run = p.add_run("Web框架：")
run.bold = True
run.font.size = Pt(12)
run = p.add_run(" FastAPI")

p = doc.add_paragraph()
run = p.add_run("API接口：")
run.bold = True
run.font.size = Pt(12)
run = p.add_run(" RESTful API with CORS support")

doc.add_paragraph()

# 核心功能
h1 = doc.add_heading("三、核心功能", level=1)

# 1. 危机检测
h2 = doc.add_heading("1. 危机检测（Crisis Detection）", level=2)
p = doc.add_paragraph()
p.add_run("• 自杀/自伤关键词检测").bold = True
p.add_run(' - 检测"不想活了"、"想死"等高危关键词')

p = doc.add_paragraph()
p.add_run("• 4级风险评估系统").bold = True
p.add_run(" - low/medium/high/critical")

p = doc.add_paragraph()
p.add_run("• 危机干预协议").bold = True
p.add_run(" - 提供专业心理援助热线")

# 2. 心理评估
h2 = doc.add_heading("2. 心理评估（Psychological Assessment）", level=2)
p = doc.add_paragraph()
p.add_run("• 焦虑自评量表（GAD-7）").bold = True
p = doc.add_paragraph()
p.add_run("• 抑郁自评量表（PHQ-9）").bold = True
p = doc.add_paragraph()
p.add_run("• 压力感知量表（PSS-10）").bold = True
p = doc.add_paragraph()
p.add_run("• 个性化建议生成").bold = True

# 3. 知识库检索
h2 = doc.add_heading("3. 知识库检索（RAG Retrieval）", level=2)
p = doc.add_paragraph()
p.add_run("• 基于关键词的语义检索").bold = True
p = doc.add_paragraph()
p.add_run("• 心理知识分类").bold = True
p.add_run(" - 心理咨询、疗愈技术、中华智慧等")
p = doc.add_paragraph()
p.add_run("• 相关性评分与排序").bold = True

# 4. 记忆系统
h2 = doc.add_heading("4. 记忆系统（Memory System）", level=2)
p = doc.add_paragraph()
p.add_run("• 短期记忆").bold = True
p.add_run(" - 会话上下文管理")
p = doc.add_paragraph()
p.add_run("• 长期记忆").bold = True
p.add_run(" - 用户画像与历史交互")
p = doc.add_paragraph()
p.add_run("• 用户特征提取").bold = True
p.add_run(" - 个性化服务支持")

# 5. 对话引擎
h2 = doc.add_heading("5. 对话引擎（Dialog Engine）", level=2)
p = doc.add_paragraph()
p.add_run("• 意图识别").bold = True
p.add_run(" - 情感支持、知识问答、寻求帮助、练习请求、一般对话")
p = doc.add_paragraph()
p.add_run("• 情感共情回应").bold = True
p = doc.add_paragraph()
p.add_run("• DeepSeek LLM 集成").bold = True
p.add_run(" - 实现智能对话生成")

doc.add_paragraph()

# 项目结构
h1 = doc.add_heading("四、项目结构", level=1)

structure = """
minghe-companion/
├── src/
│   ├── agents/
│   │   └── psychology_master.py    # 心理资讯大师Agent
│   ├── llm/
│   │   └── client.py               # DeepSeek LLM客户端
│   ├── tools/
│   │   ├── crisis.py                # 危机检测
│   │   ├── rag.py                  # 知识检索
│   │   └── assessment.py           # 心理评估
│   ├── memory/
│   │   └── system.py               # 记忆系统
│   ├── api/
│   │   └── main.py                 # FastAPI应用
│   └── core/
│       ├── config.py                # 配置管理
│       ├── prompt.py                # 提示词模板
│       └── constants.py             # 常量定义
├── tests/                           # 测试套件
│   ├── test_tools/
│   │   ├── test_crisis.py          # 10个测试
│   │   ├── test_rag.py              # 18个测试
│   │   └── test_assessment.py      # 23个测试
│   └── test_memory/
│       └── test_system.py           # 33个测试
├── knowledge_base/                  # 知识库
├── frontend.html                    # 前端界面
├── .env                            # 环境配置
└── pyproject.toml                  # 项目配置
"""
p = doc.add_paragraph(structure)
for run in p.runs:
    run.font.size = Pt(9)
    run.font.name = "Consolas"

doc.add_paragraph()

# 测试覆盖
h1 = doc.add_heading("五、测试覆盖", level=1)

p = doc.add_paragraph()
run = p.add_run("项目包含 ")
run.font.size = Pt(12)
run = p.add_run("84 个单元测试")
run.bold = True
run.font.size = Pt(12)
run = p.add_run("，覆盖所有核心模块：")
run.font.size = Pt(12)

# 测试表格
table = doc.add_table(rows=5, cols=3)
table.style = "Light Grid Accent 1"

# 表头
header_cells = table.rows[0].cells
header_cells[0].text = "测试类别"
header_cells[1].text = "测试数量"
header_cells[2].text = "状态"

# 数据行
data = [
    ["危机检测", "10", "✅ 通过"],
    ["RAG检索", "18", "✅ 通过"],
    ["心理评估", "23", "✅ 通过"],
    ["记忆系统", "33", "✅ 通过"],
]

for i, row_data in enumerate(data):
    cells = table.rows[i + 1].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()

# API 端点
h1 = doc.add_heading("六、API 端点", level=1)

p = doc.add_paragraph()
p.add_run("• GET /health ").bold = True
p.add_run("- 健康检查")

p = doc.add_paragraph()
p.add_run("• POST /chat ").bold = True
p.add_run("- 与Agent对话")

p = doc.add_paragraph()
p.add_run("• GET /assessment/template/{type} ").bold = True
p.add_run("- 获取评估问卷模板")

p = doc.add_paragraph()
p.add_run("• POST /assessment ").bold = True
p.add_run("- 提交心理评估")

p = doc.add_paragraph()
p.add_run("• GET /user/{user_id}/profile ").bold = True
p.add_run("- 获取用户画像")

doc.add_paragraph()

# 使用说明
h1 = doc.add_heading("七、使用说明", level=1)

p = doc.add_paragraph()
p.add_run("1. 启动API服务").bold = True

p = doc.add_paragraph()
run = p.add_run("   cd minghe-companion")
run.font.name = "Consolas"
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run("   uvicorn src.api.main:app --port 8000")
run.font.name = "Consolas"
run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run("2. 打开前端界面").bold = True

p = doc.add_paragraph()
run = p.add_run("   使用浏览器打开 frontend.html")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.add_run("3. 配置 DeepSeek API Key").bold = True

p = doc.add_paragraph()
run = p.add_run("   在 .env 文件中设置 DEEPSEEK_API_KEY")
run.font.size = Pt(12)

doc.add_paragraph()

# 总结
h1 = doc.add_heading("八、项目总结", level=1)

p = doc.add_paragraph()
p.add_run("明禾陪伴项目已完成以下工作：").bold = True

p = doc.add_paragraph()
p.add_run("✅ 完整的心理健康 AI Agent 系统").bold = True

p = doc.add_paragraph()
p.add_run("✅ 84个单元测试，全部通过").bold = True

p = doc.add_paragraph()
p.add_run("✅ DeepSeek 大语言模型集成").bold = True

p = doc.add_paragraph()
p.add_run("✅ FastAPI RESTful 接口").bold = True

p = doc.add_paragraph()
p.add_run("✅ 响应式前端界面").bold = True

p = doc.add_paragraph()
p.add_run("✅ 完整的危机检测与干预机制").bold = True

doc.add_paragraph()

# 页脚
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
run.font.size = Pt(10)

# 保存文档
output_path = "C:/Users/qqnag/Desktop/明禾陪伴项目完成报告.docx"
doc.save(output_path)
print(f"报告已保存到：{output_path}")
